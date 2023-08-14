import pyttsx3
import pywhatkit
import speech_recognition as sr
import webbrowser
from AppOpener import open
import  wikipedia
import docx
import random
engine = pyttsx3.init()

r = sr.Recognizer()
with sr.Microphone() as source:
    r.adjust_for_ambient_noise(source, duration=2)
    text = "Hello, I am your Assistant How can I help you today"
    print(text)
    engine.say(text)
    engine.runAndWait()
    audio2 = r.listen(source)
    Ques = r.recognize_google(audio2)
    Ques = Ques.lower()
    
if Ques == "take a note":
    doc = docx.Document()
    engine.say("Please enter your note")
    engine.runAndWait()
    note = input("Please enter your note")
    doc.add_paragraph(note)
    qwerty = random.randrange(1,1000)
    doc.save("C:/Your Note",qwerty")
             
if Ques == "what can you do":
    ans1 = "i can search for you and answer your queries"
    engine.say(ans1)
    engine.runAndWait()

if Ques == "who am i":
    ans01 = "You are an amazing youtuber"
    engine.say(ans01)
    engine.runAndWait()

if Ques == "open google":
    webbrowser.open("www.google.com")

if Ques == "open whatsapp":
    webbrowser.open("https://web.whatsapp.com/")

if Ques == "open stack":
    webbrowser.open("https://stackoverflow.com/")


if Ques == "open app":
    app_call = "which app"
    engine.say(app_call)
    engine.runAndWait()
    app = input("Which app:")
    engine.say("opening", app)
    engine.runAndWait()
    open(app)
a = Ques.split()

if any(a) == "write":
    n = wikipedia.summary(Ques[12:])
    print(n)

for a in a:
    if a == "whatsapp":
        engine.say("Enter the number of the person you wanna message")
        engine.runAndWait()
        Number_of_person = input("Enter the number of the person you wanna message")
        engine.say("Enter the message")
        engine.runAndWait()
        Message_what = "Enter the message"
        pywhatkit.sendwhatmsg_instantly(Number_of_person,Message_what)

if Ques == "Take a screenshot":
    pywhatkit.take_screenshot()

if Ques == "make that weird art":
    engine.say("enter image name")
    engine.runAndWait()
    image = "enter image name:"
    engine.say("Enter text file name")
    text_file = "enter text file name:"
    pywhatkit.image_to_ascii_art(image,text_file)

else:
    pywhatkit.search(Ques)

with sr.Microphone() as source:
    r.adjust_for_ambient_noise(source, duration=2)
    text1 = "anything else  can help you with"
    print(text1)
    engine.say(text1)
    engine.runAndWait()
    audio2 = r.listen(source)
    Ques2 = r.recognize_google(audio2)
    Ques2 = Ques2.lower()

if Ques2 == "no":
    engine.say("Okay, bye!")
    engine.runAndWait()


if Ques2 == "what can you do":
    ans1 = "i can search for you and answer your queries"
    engine.say(ans1)
    engine.runAndWait()

if Ques2 == "who am i":
    ans01 = "You are an amazing youtuber"
    engine.say(ans01)
    engine.runAndWait()

if Ques2 == "open google":
    webbrowser.open("www.google.com")

if Ques2 == "open whatsapp":
    webbrowser.open("https://web.whatsapp.com/")

if Ques2 == "open stack":
    webbrowser.open("https://stackoverflow.com/")


if Ques2 == "open app":
    app_call = "which app"
    engine.say(app_call)
    engine.runAndWait()
    app = input("Which app:")
    engine.say("opening", app)
    engine.runAndWait()
    open(app)


b = Ques2.split()
for a in b:
    if a == "whatsapp":
        engine.say("Enter the number of the person you wanna message")
        engine.runAndWait()
        Number_of_person = input("Enter the number of the person you wanna message")
        engine.say("Enter the message")
        engine.runAndWait()
        Message_what = "Enter the message"
        pywhatkit.sendwhatmsg_instantly(Number_of_person,Message_what)

if Ques2 == "Take a screenshot":
    pywhatkit.take_screenshot()

if Ques2 == "make that weird art":
    engine.say("enter image name")
    engine.runAndWait()
    image = "enter image name:"
    engine.say("Enter text file name")
    text_file = "enter text file name:"
    pywhatkit.image_to_ascii_art(image,text_file)

if any(b) == "write":
    n = wikipedia.summary(Ques2[12:])
    print(n)
if Ques2 == "take a note":
    doc = docx.Document()
    engine.say("Please enter your note")
    engine.runAndWait()
    note = input("Please enter your note")
    doc.add_paragraph(note)
    qwerty = random.randrange(1,1000)
    doc.save("C:/Your Note",qwerty")
else:
    pywhatkit.search(Ques2)

    





